import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import google.generativeai as genai
import io
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import numpy as np

# ==========================================
# 0. CẤU HÌNH HỆ THỐNG & GIAO DIỆN
# ==========================================
st.set_page_config(page_title="Quản trị KHTN 2026 - THPT Yên Minh", page_icon="🎓", layout="wide")

# CSS để làm đẹp giao diện
st.markdown("""
    <style>
    .main { background-color: #f5f7f9; }
    .stButton>button { width: 100%; border-radius: 5px; height: 3em; font-weight: bold; }
    .report-header { text-align: center; color: #1A365D; font-weight: 900; text-transform: uppercase; }
    </style>
    """, unsafe_allow_html=True)

st.markdown("<h1 class='report-header'>HỆ SINH THÁI QUẢN TRỊ CHẤT LƯỢNG ĐA CHIỀU - THPT YÊN MINH</h1>", unsafe_allow_html=True)
st.markdown("<hr style='margin-bottom: 30px;'>", unsafe_allow_html=True)

# ==========================================
# 1. HÀM XUẤT WORD CHUẨN NGHỊ ĐỊNH 30
# ==========================================
def tao_file_word_chuan_nd30(noi_dung_ai, tieu_de_van_ban):
    doc = docx.Document()
    # Cấu hình lề trang chuẩn: Trên 2cm, Dưới 2cm, Trái 3cm, Phải 1.5cm
    section = doc.sections[0]
    section.top_margin, section.bottom_margin = Cm(2), Cm(2)
    section.left_margin, section.right_margin = Cm(3), Cm(1.5)

    # Header: Quốc hiệu & Tên trường
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Cm(7.5)
    table.columns[1].width = Cm(9.5)
    
    # Cột trái: Tên đơn vị
    c1 = table.cell(0, 0).paragraphs[0]
    r1 = c1.add_run("TRƯỜNG THPT YÊN MINH\n")
    r1.font.size = Pt(12)
    r1.bold = True
    r2 = c1.add_run("TỔ TỰ NHIÊN")
    r2.font.size = Pt(12)
    r2.bold = True
    c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Cột phải: Quốc hiệu
    c2 = table.cell(0, 1).paragraphs[0]
    r3 = c2.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    r3.font.size = Pt(12)
    r3.bold = True
    r4 = c2.add_run("Độc lập - Tự do - Hạnh phúc")
    r4.font.size = Pt(13)
    r4.bold = True
    c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tiêu đề báo cáo
    doc.add_paragraph("\n")
    p_td = doc.add_paragraph()
    r_td = p_td.add_run(tieu_de_van_ban.upper())
    r_td.font.size = Pt(14)
    r_td.bold = True
    p_td.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Nội dung
    for line in noi_dung_ai.split('\n'):
        line = line.strip()
        if line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
                if idx % 2 != 0: run.bold = True
            if line.startswith(('*', '-')): p.style = 'List Bullet'

    # Phần ký tên
    doc.add_paragraph("\n")
    sign_table = doc.add_table(rows=1, cols=2)
    nn = sign_table.cell(0, 0).paragraphs[0]
    nn.add_run("Nơi nhận:\n- BGH (để b/c);\n- Lưu VT, Tổ chuyên môn.").font.size = Pt(11)
    
    nk = sign_table.cell(0, 1).paragraphs[0]
    nk.add_run("PHÓ HIỆU TRƯỞNG\n").bold = True
    nk.add_run("(Phụ trách Chuyên môn)").italic = True
    nk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ==========================================
# 2. XỬ LÝ DỮ LIỆU LÕI (CLEAN DATA 100%)
# ==========================================
@st.cache_data(ttl=10)
def load_and_transform_data(url):
    try:
        file_id = url.split("/d/")[1].split("/")[0]
        export_url = f"https://docs.google.com/spreadsheets/d/{file_id}/export?format=csv"
        df_raw = pd.read_csv(export_url, header=None)
        
        # Tìm dòng tiêu đề môn học
        idx_mon = -1
        for i in range(min(10, len(df_raw))):
            row_str = " ".join([str(x).lower() for x in df_raw.iloc[i].values])
            if "toán" in row_str and "văn" in row_str:
                idx_mon = i; break
        if idx_mon == -1: idx_mon = 1
        
        row_lan = df_raw.iloc[idx_mon-1].copy()
        row_mon = df_raw.iloc[idx_mon].copy()
        
        # Đọc chỉ tiêu tự động
        ct_chung = 6.0; dict_ct_mon = {}
        for i in range(min(20, len(df_raw))):
            row_str = " ".join([str(x).lower() for x in df_raw.iloc[i].values])
            if "chỉ tiêu chung" in row_str:
                m = re.search(r'(\d+[\.,]\d+|\d+)', row_str)
                if m: ct_chung = float(m.group().replace(',', '.'))
            if "chỉ tiêu" in row_str:
                for j, val in enumerate(df_raw.iloc[i].values):
                    mon = str(row_mon.iloc[j]).strip()
                    if mon and j > 3:
                        try: dict_ct_mon[mon] = float(str(val).replace(',', '.'))
                        except: pass

        # Transform dữ liệu
        active_lan = "Lần 1"; new_cols = []
        for i in range(len(row_mon)):
            c_lan = str(row_lan.iloc[i]).strip()
            if "Lần" in c_lan: active_lan = c_lan
            c_mon = str(row_mon.iloc[i]).strip()
            if i <= 3: new_cols.append(c_mon)
            else: new_cols.append(f"{c_mon}|{active_lan}")
            
        df_ngang = df_raw.iloc[idx_mon+1:].reset_index(drop=True)
        df_ngang.columns = new_cols
        
        # Melting & Cleaning
        df_doc = pd.melt(df_ngang, id_vars=new_cols[:4], var_name='Split', value_name='Diem_Thi')
        df_doc[['Mon_Hoc', 'Lan_Thi']] = df_doc['Split'].str.split('|', expand=True)
        
        # BỘ LỌC THÔNG MINH: Loại bỏ dòng rác & Học sinh vắng
        def filter_real_student(row):
            name = str(row.iloc[1]).lower() # Cột Họ tên
            lop = str(row.iloc[2]).lower() # Cột Lớp
            if any(k in name or k in lop for k in ['chỉ tiêu', 'trung bình', 'tổng', 'tỉ lệ', 'đtb']): return False
            if name in ['', 'nan', 'none']: return False
            return True
            
        df_doc = df_doc[df_doc.apply(filter_real_student, axis=1)]
        df_doc['Diem_Thi'] = pd.to_numeric(df_doc['Diem_Thi'].astype(str).str.replace(',', '.'), errors='coerce')
        
        # QUAN TRỌNG: Loại bỏ thí sinh không dự thi (NaN) để điểm TB chính xác
        df_doc = df_doc.dropna(subset=['Diem_Thi'])
        
        # Rename cột chuẩn
        df_doc = df_doc.rename(columns={new_cols[1]: 'Ten_Hoc_Sinh', new_cols[2]: 'Lop'})
        return df_doc, sorted(df_doc['Lop'].unique()), ct_chung, dict_ct_mon, None
    except Exception as e:
        return None, None, 6.0, {}, str(e)

# ==========================================
# 3. GIAO DIỆN & TÁC VỤ AI
# ==========================================
with st.sidebar:
    st.header("⚙️ QUẢN TRỊ")
    admin_pw = st.text_input("Mật khẩu:", type="password")
    is_admin = (admin_pw == st.secrets.get("ADMIN_PASSWORD", "123"))
    gsheet_url = st.text_input("Link Google Sheet:")

if gsheet_url:
    df, ds_lop, ct_chung, ct_mon_dict, err = load_and_transform_data(gsheet_url)
    if err: st.error(f"Lỗi: {err}")
    else:
        ds_lan = sorted(df['Lan_Thi'].unique())
        ds_mon = sorted(df['Mon_Hoc'].unique())
        
        tab1, tab2, tab3, tab4, tab5 = st.tabs(["📊 TỔNG QUAN", "📈 TIẾN TRÌNH", "🔎 PHÂN TÍCH", "🏆 TỔNG HỢP", "🎓 XÉT TUYỂN"])
        
        # --- TAB 1: TỔNG QUAN ---
        with tab1:
            st.subheader("🌟 Biến động chất lượng toàn khóa")
            df_tb = df.groupby('Lan_Thi')['Diem_Thi'].mean().reset_index()
            fig = px.line(df_tb, x='Lan_Thi', y='Diem_Thi', markers=True, text=df_tb['Diem_Thi'].round(2))
            fig.add_hline(y=ct_chung, line_dash="dash", line_color="red", annotation_text="Chỉ tiêu")
            st.plotly_chart(fig, use_container_width=True)
            st.dataframe(df_tb.rename(columns={'Diem_Thi': 'Điểm TB'}), use_container_width=True)

        # --- TAB 3: PHÂN TÍCH (CÓ XUẤT WORD) ---
        with tab3:
            col1, col2 = st.columns(2)
            m_sel = col1.selectbox("Môn:", ds_mon)
            l_sel = col2.selectbox("Đợt:", ds_lan)
            
            ct_m = ct_mon_dict.get(m_sel, 6.5)
            df_m = df[(df['Mon_Hoc'] == m_sel) & (df['Lan_Thi'] == l_sel)]
            
            # Tính toán bảng xếp hạng
            df_rank = df_m.groupby('Lop')['Diem_Thi'].agg(['mean', 'count']).reset_index()
            df_rank.columns = ['Lớp', 'Điểm TB', 'Sĩ số']
            df_rank['Chênh lệch'] = (df_rank['Điểm TB'] - ct_m).round(2)
            df_rank = df_rank.sort_values('Điểm TB', ascending=False)
            
            # Dòng toàn khối
            tb_k = df_m['Diem_Thi'].mean()
            df_rank.loc[len(df_rank)] = ['⭐ TOÀN KHỐI', tb_k, len(df_m), round(tb_k - ct_m, 2)]
            
            st.dataframe(df_rank, use_container_width=True, hide_index=True)
            
            if st.button("🤖 AI PHÂN TÍCH & THAM MƯU", key="ai3"):
                with st.spinner("AI đang soạn thảo báo cáo..."):
                    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
                    model = genai.GenerativeModel('gemini-1.5-flash')
                    prompt = f"Phân tích kết quả thi môn {m_sel} đợt {l_sel}. Chỉ tiêu {ct_m}. Số liệu: {df_rank.to_string()}. Đưa ra đánh giá và 3 giải pháp sư phạm thực chiến."
                    res = model.generate_content(prompt).text
                    st.session_state.rep3 = res
                    st.info(res)
            
            if "rep3" in st.session_state:
                word_file = tao_file_word_chuan_nd30(st.session_state.rep3, f"BÁO CÁO PHÂN TÍCH CHUYÊN MÔN MÔN {m_sel.upper()}")
                st.download_button("📄 Tải Báo cáo Word (Trình ký)", word_file, f"Bao_cao_{m_sel}.docx", use_container_width=True)

        # --- TAB 4: TỔNG HỢP ---
        with tab4:
            st.subheader("🏆 So sánh tiến độ giữa 2 đợt thi")
            l1 = st.selectbox("Đợt trước:", ds_lan, index=0)
            l2 = st.selectbox("Đợt sau:", ds_lan, index=len(ds_lan)-1)
            
            # Logic tính toán chênh lệch (Pivot table)
            df_comp = df[df['Lan_Thi'].isin([l1, l2])].groupby(['Lop', 'Lan_Thi'])['Diem_Thi'].mean().unstack().reset_index()
            df_comp['Thay đổi'] = (df_comp[l2] - df_comp[l1]).round(2)
            st.dataframe(df_comp, use_container_width=True)
            
            if st.button("🤖 AI PHÂN TÍCH TOÀN TRƯỜNG", key="ai4"):
                genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
                model = genai.GenerativeModel('gemini-1.5-flash')
                prompt = f"Phân tích sự tiến bộ toàn trường từ {l1} đến {l2}. Số liệu: {df_comp.to_string()}. Vinh danh lớp khá, cảnh báo lớp yếu và đề xuất giải pháp quản trị cho Phó Hiệu trưởng."
                res = model.generate_content(prompt).text
                st.session_state.rep4 = res
                st.info(res)
                
            if "rep4" in st.session_state:
                word_file = tao_file_word_chuan_nd30(st.session_state.rep4, "BÁO CÁO TỔNG HỢP TIẾN ĐỘ CHẤT LƯỢNG TOÀN TRƯỜNG")
                st.download_button("📄 Tải Báo cáo Word (Trình ký)", word_file, "Bao_cao_Toan_Truong.docx", use_container_width=True)

        # --- TAB 5: XÉT TUYỂN ---
        with tab5:
            st.subheader("🎓 Dự toán Tổ hợp Đại học 2026")
            # Tự động tính các tổ hợp phổ biến
            df_wide = df[df['Lan_Thi'] == ds_lan[-1]].pivot_table(index=['Ten_Hoc_Sinh', 'Lop'], columns='Mon_Hoc', values='Diem_Thi').reset_index()
            
            # Hàm tính điểm tổ hợp an toàn
            def calc(r, subjects): 
                try: return round(sum(r[s] for s in subjects), 2)
                except: return None

            if 'Toán' in df_wide and 'Lý' in df_wide and 'Hóa' in df_wide:
                df_wide['A00'] = df_wide.apply(lambda r: calc(r, ['Toán', 'Lý', 'Hóa']), axis=1)
            if 'Văn' in df_wide and 'Sử' in df_wide and 'Địa' in df_wide:
                df_wide['C00'] = df_wide.apply(lambda r: calc(r, ['Văn', 'Sử', 'Địa']), axis=1)
            if 'Toán' in df_wide and 'Văn' in df_wide and 'Anh' in df_wide:
                df_wide['D01'] = df_wide.apply(lambda r: calc(r, ['Toán', 'Văn', 'Anh']), axis=1)
                
            st.dataframe(df_wide, use_container_width=True)